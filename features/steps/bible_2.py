import time

import docx
from behave import *

import oR
from Selenium_Operations.Element_Operations import Element_Operations

use_step_matcher("parse")

@given(u'User navigated to Bible and write data')
def step_impl(context):
    verse = 1
    chapter = 1
    while chapter <= 8:
        context.ele_ops = Element_Operations(context.driver)
        verse_str = str(verse)
        chapter_str = str(chapter)
        context.ele_ops.get_url("https://www.biblehub.com/songs/"+chapter_str+"-"+verse_str+".htm")
        context.ele_ops.maximize_window()
        if context.ele_ops.get_element_text(oR.top_heading__XPATH) != "Bible Hub":
            verse_text = context.ele_ops.get_text_from_multiple_elements(oR.bible_text_verse__XPATH)
            mydoc = docx.Document("E://auto_frameworks//Behave_python//features//steps//songs.docx")
            p = mydoc.add_paragraph()
            super_text = p.add_run(chapter_str + ":" + verse_str)
            super_text.font.superscript = True
            p.add_run([" " + ele + " " for ele in verse_text])
            mydoc.save("E://auto_frameworks//Behave_python//features//steps//songs.docx")
        elif context.ele_ops.get_element_text(oR.top_heading__XPATH) == "Bible Hub":
            verse = 1
            verse -= 1
            chapter += 1
        verse += 1


